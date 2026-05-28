from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Estilos ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ---- Portada ----
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Guía de Despliegue\nProyecto Asistecia')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Spring Boot 4.0.6 + Java 17 + MySQL\nServidor Ubuntu')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ---- Tabla de Contenido ----
doc.add_heading('Índice', level=1)
toc_items = [
    '1. Requisitos Previos',
    '2. Instalación de Java 17',
    '3. Instalación y Configuración de MySQL',
    '4. Subir el Proyecto al Servidor',
    '5. Compilar y Empaquetar con Maven',
    '6. Ejecutar la Aplicación',
    '7. Configurar como Servicio systemd',
    '8. Variables de Entorno',
    '9. Script de Automatización (deploy.sh)',
    '10. Verificación y Solución de Problemas',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
def add_step(heading, steps, code_blocks=None):
    doc.add_heading(heading, level=1)
    for s in steps:
        p = doc.add_paragraph(s, style='List Number')
    if code_blocks:
        for cb in code_blocks:
            doc.add_paragraph()
            code_p = doc.add_paragraph()
            code_p.paragraph_format.left_indent = Cm(1)
            run = code_p.add_run(cb)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            run.font.highlight_color = RGBColor(0xF5, 0xF5, 0xF5)  # not directly supported, use shading
            # Add shading via XML
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            shading_elm.set(qn('w:val'), 'clear')
            code_p.runs[0]._element.get_or_add_rPr().append(shading_elm)

# ---- 1 ----
doc.add_heading('1. Requisitos Previos', level=1)
doc.add_paragraph('Antes de comenzar, asegúrate de tener acceso a:')
doc.add_paragraph('Servidor Ubuntu 20.04 o superior (con acceso SSH)', style='List Bullet')
doc.add_paragraph('Usuario con permisos sudo', style='List Bullet')
doc.add_paragraph('Conexión a Internet para descargar dependencias', style='List Bullet')
doc.add_paragraph('Los archivos del proyecto (carpeta asistecia)', style='List Bullet')

# ---- 2 ----
doc.add_heading('2. Instalación de Java 17', level=1)
doc.add_paragraph('El proyecto requiere Java 17 (OpenJDK). Ejecuta los siguientes comandos:')

code = 'sudo apt update\nsudo apt install openjdk-17-jdk -y\njava -version'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Verifica que la salida muestre "openjdk version 17.x.x".')

# ---- 3 ----
doc.add_heading('3. Instalación y Configuración de MySQL', level=1)

doc.add_heading('3.1 Instalar MySQL', level=2)
code = 'sudo apt install mysql-server -y\nsudo mysql_secure_installation'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('3.2 Crear Base de Datos y Usuario', level=2)

p = doc.add_paragraph()
run = p.add_run('Ingresa a la consola de MySQL:\n')
run.font.size = Pt(11)
code = 'sudo mysql -u root -p'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Dentro de MySQL, ejecuta:')
code = """CREATE DATABASE gescol;
CREATE USER 'root'@'localhost' IDENTIFIED BY '1651';
GRANT ALL PRIVILEGES ON gescol.* TO 'root'@'localhost';
FLUSH PRIVILEGES;
EXIT;"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Nota: La contraseña y usuario deben coincidir con los definidos en src/main/resources/application.properties.', style='List Bullet')
doc.add_paragraph('Para entornos productivos, cambia la contraseña por una más segura.', style='List Bullet')

# ---- 4 ----
doc.add_heading('4. Subir el Proyecto al Servidor', level=1)

doc.add_heading('Opción A: SCP (Recomendada)', level=2)
doc.add_paragraph('Desde tu máquina local (Windows con PowerShell):')
code = 'scp -r C:\\ruta\\a\\asistecia usuario@tu-servidor:/home/usuario/'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Opción B: Git', level=2)
doc.add_paragraph('Si el proyecto está en GitHub/GitLab:')
code = 'git clone https://github.com/tu-usuario/asistecia.git'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Opción C: SFTP con FileZilla / WinSCP', level=2)
doc.add_paragraph('Usa un cliente gráfico SFTP para transferir la carpeta completa del proyecto.')

# ---- 5 ----
doc.add_heading('5. Compilar y Empaquetar con Maven', level=1)
doc.add_paragraph('Accede al directorio del proyecto y compila:')

code = 'cd /home/usuario/asistecia\nchmod +x mvnw\n./mvnw clean package -DskipTests'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Si el comando falla por falta de memoria, usa:')
code = './mvnw clean package -DskipTests -Dmaven.test.skip=true -Xmx512m'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Al finalizar, el archivo JAR se generará en:')
p = doc.add_paragraph()
run = p.add_run('target/asistecia-0.0.1-SNAPSHOT.jar')
run.bold = True

# ---- 6 ----
doc.add_heading('6. Ejecutar la Aplicación', level=1)

doc.add_heading('6.1 Ejecución Directa', level=2)
code = 'java -jar target/asistecia-0.0.1-SNAPSHOT.jar'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('La aplicación iniciará en el puerto 8080. Accede en: http://tu-servidor:8080')

doc.add_heading('6.2 Ejecución en Segundo Plano', level=2)
code = 'nohup java -jar target/asistecia-0.0.1-SNAPSHOT.jar > app.log 2>&1 &'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Para ver los logs en vivo:')
code = 'tail -f app.log'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ---- 7 ----
doc.add_heading('7. Configurar como Servicio systemd', level=1)
doc.add_paragraph('Esto permite que la aplicación inicie automáticamente al reiniciar el servidor.')

doc.add_heading('7.1 Crear el archivo de servicio', level=2)
code = 'sudo nano /etc/systemd/system/asistecia.service'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Pega el siguiente contenido:')
code = """[Unit]
Description=Asistecia Spring Boot App
After=network.target mysql.service

[Service]
User=ubuntu
WorkingDirectory=/home/ubuntu/asistecia
ExecStart=/usr/bin/java -jar /home/ubuntu/asistecia/target/asistecia-0.0.1-SNAPSHOT.jar
SuccessExitStatus=143
Restart=always
RestartSec=10

[Install]
WantedBy=multi-user.target"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Ajusta "User" y las rutas según tu configuración.')

doc.add_heading('7.2 Habilitar e iniciar el servicio', level=2)
code = """sudo systemctl daemon-reload
sudo systemctl enable asistecia
sudo systemctl start asistecia
sudo systemctl status asistecia"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ---- 8 ----
doc.add_heading('8. Variables de Entorno', level=1)
doc.add_paragraph('Para no hardcodear credenciales en application.properties, usa variables de entorno:')

code = """export SPRING_DATASOURCE_URL=jdbc:mysql://localhost:3306/gescol
export SPRING_DATASOURCE_USERNAME=root
export SPRING_DATASOURCE_PASSWORD=1651
java -jar target/asistecia-0.0.1-SNAPSHOT.jar"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Para systemd, agrega las variables en la sección [Service]:')
code = """[Service]
Environment="SPRING_DATASOURCE_PASSWORD=1651"
Environment="SPRING_DATASOURCE_URL=jdbc:mysql://localhost:3306/gescol\""""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ---- 9 ----
doc.add_heading('9. Script de Automatización (deploy.sh)', level=1)

doc.add_paragraph('Para automatizar todo el proceso, se incluye el script deploy.sh. Este script realiza todos los pasos anteriores de forma automática.')

doc.add_heading('9.1 Requisitos', level=2)
doc.add_paragraph('El proyecto debe estar copiado en /home/tu-usuario/asistecia antes de ejecutar el script.', style='List Bullet')
doc.add_paragraph('Ejecutar con permisos de superusuario: sudo bash deploy.sh', style='List Bullet')

doc.add_heading('9.2 Contenido del script', level=2)
code = '''#!/bin/bash
# ===================================================
# Script de despliegue automático - Proyecto Asistecia
# Spring Boot 4.0.6 + Java 17 + MySQL
# Uso: sudo bash deploy.sh
# ===================================================

set -e

APP_NAME="asistecia"
APP_DIR="/home/$SUDO_USER/$APP_NAME"
JAR_NAME="$APP_NAME-0.0.1-SNAPSHOT.jar"
DB_NAME="gescol"
DB_USER="root"
DB_PASS="1651"
SPRING_PORT="8080"

echo "========================================"
echo "  Despliegue Automático - $APP_NAME"
echo "========================================"

# 1. Actualizar sistema
echo "[1/7] Actualizando paquetes del sistema..."
apt update && apt upgrade -y

# 2. Instalar Java 17
echo "[2/7] Instalando Java 17..."
if ! java -version 2>&1 | grep -q 'openjdk version "17'; then
    apt install openjdk-17-jdk -y
fi

# 3. Instalar MySQL
echo "[3/7] Instalando MySQL..."
if ! command -v mysql &> /dev/null; then
    apt install mysql-server -y
fi

# 4. Configurar base de datos
echo "[4/7] Configurando base de datos $DB_NAME..."
mysql -u root <<SQL
CREATE DATABASE IF NOT EXISTS $DB_NAME;
ALTER USER '$DB_USER'@'localhost' IDENTIFIED BY '$DB_PASS';
GRANT ALL PRIVILEGES ON $DB_NAME.* TO '$DB_USER'@'localhost';
FLUSH PRIVILEGES;
SQL

# 5. Compilar proyecto
echo "[5/7] Compilando proyecto con Maven..."
cd "$APP_DIR"
chmod +x mvnw
./mvnw clean package -DskipTests -q

# 6. Crear servicio systemd
echo "[6/7] Configurando servicio systemd..."
SERVICE_FILE="/etc/systemd/system/$APP_NAME.service"
cat > "$SERVICE_FILE" <<EOF
[Unit]
Description=$APP_NAME Spring Boot App
After=network.target mysql.service

[Service]
User=$SUDO_USER
WorkingDirectory=$APP_DIR
ExecStart=/usr/bin/java -jar $APP_DIR/target/$JAR_NAME
SuccessExitStatus=143
Restart=always
RestartSec=10
Environment=SPRING_DATASOURCE_URL=jdbc:mysql://localhost:3306/$DB_NAME
Environment=SPRING_DATASOURCE_USERNAME=$DB_USER
Environment=SPRING_DATASOURCE_PASSWORD=$DB_PASS

[Install]
WantedBy=multi-user.target
EOF
systemctl daemon-reload
systemctl enable "$APP_NAME"
systemctl start "$APP_NAME"

# 7. Verificar
echo "[7/7] Verificando despliegue..."
sleep 5
if systemctl is-active --quiet "$APP_NAME"; then
    echo "Despliegue completado con EXITO!"
    echo "App: http://$(curl -s ifconfig.me):$SPRING_PORT"
else
    echo "ERROR: El servicio no inició. Revisa: sudo journalctl -u $APP_NAME -f"
fi'''

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_heading('9.3 Cómo usar el script', level=2)
steps_script = [
    'Sube el proyecto al servidor en /home/tu-usuario/asistecia',
    'Sube también el archivo deploy.sh al mismo nivel',
    'Da permisos de ejecución: chmod +x deploy.sh',
    'Ejecuta: sudo bash deploy.sh',
]
for s in steps_script:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()
doc.add_paragraph('El script tarda entre 2-5 minutos dependiendo de la velocidad del servidor. Al finalizar, la aplicación quedará corriendo como servicio systemd.', style='List Bullet')

# ---- 10 ----
doc.add_heading('10. Verificación y Solución de Problemas', level=1)

doc.add_heading('10.1 Verificar que la app está corriendo', level=2)
code = 'curl http://localhost:8080\nsudo systemctl status asistecia\nsudo journalctl -u asistecia -f'
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('10.2 Problemas comunes', level=2)

problems = [
    ('Error: "Port 8080 already in use"',
     'Usa otro puerto: java -jar app.jar --server.port=8081'),
    ('Error de conexión a MySQL',
     'Verifica que MySQL esté corriendo: sudo systemctl status mysql'),
    ('Error: "Java not found"',
     'Verifica la instalación: java -version y que JAVA_HOME esté configurado'),
    ('Error de permisos en mvnw',
     'Ejecuta: chmod +x mvnw'),
    ('Error de memoria al compilar',
     'Usa: ./mvnw clean package -DskipTests -Xmx512m'),
]

for title, solution in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run = p.add_run(solution)

# ---- Guardar ----
output_path = 'D:\\asisteciafull\\asistecia\\Guia_Despliegue_Asistecia_Ubuntu.docx'
doc.save(output_path)
print(f'Documento creado: {output_path}')
